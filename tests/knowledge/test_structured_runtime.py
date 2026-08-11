from __future__ import annotations

import json

import pytest

from nanobot.agent.tools.context import RequestContext, request_context
from nanobot.agent.tools.knowledge import (
    KnowledgePublishTool,
    KnowledgeSearchTool,
    _load_ir_draft,
)
from nanobot.knowledge.compiler import parse_frontmatter
from nanobot.knowledge.ingest import (
    read_bounded_source,
    read_docx_text_tables,
    read_image_vision_ocr,
    read_markdown_lines,
)
from nanobot.knowledge.models import KnowledgeClaim, KnowledgeEvidence, KnowledgeIR
from nanobot.knowledge.service import KnowledgeService
from nanobot.knowledge.store import KnowledgeStore
from nanobot.security.workspace_access import (
    bind_workspace_scope,
    build_workspace_scope,
    reset_workspace_scope,
)
from nanobot.session.interaction_state import resolve_interaction
from nanobot.session.manager import SessionManager


def test_ir_round_trip_preserves_claim_evidence_and_legacy_fields() -> None:
    evidence = KnowledgeEvidence(
        source_path="raw/sources/paper.md",
        start_line=4,
        end_line=6,
        quote="bounded observation",
        extraction_method="markdown_lines",
        confidence=0.91,
    )
    ir = KnowledgeIR(
        project_id="project",
        source_path="paper.md",
        claims=[KnowledgeClaim.from_dict({
            "subject": "Runtime",
            "predicate": "supports",
            "object": "tool calls",
            "evidence": [evidence.to_dict()],
            "source_path": "paper.md",
            "confidence": 0.8,
        })],
        evidence=[evidence],
        review_hints=[{"kind": "confirmation", "title": "Check claim"}],
        relation_confidence={"runtime->tools": 0.7},
    )

    restored = KnowledgeIR.from_dict(ir.to_dict())

    assert restored.source_path == "paper.md"
    assert restored.claims[0].predicate == "supports"
    assert restored.claims[0].evidence[0].start_line == 4
    assert restored.evidence[0].quote == "bounded observation"
    assert restored.review_hints[0]["kind"] == "confirmation"
    assert restored.relation_confidence["runtime->tools"] == 0.7
    # A pre-claims IR remains readable and receives empty extension fields.
    legacy = KnowledgeIR.from_dict({"project_id": "legacy", "source_path": "old.md"})
    assert legacy.claims == []
    assert legacy.evidence == []


def test_large_ir_draft_is_loaded_only_from_project_root(tmp_path) -> None:
    store = KnowledgeStore(tmp_path)
    project_root = store.project_path("project")
    draft = project_root / "knowledge" / "drafts" / "source.json"
    draft.parent.mkdir(parents=True)
    draft.write_text(
        json.dumps({"project_id": "project", "source_path": "source.doc", "pages": []}),
        encoding="utf-8",
    )

    loaded = _load_ir_draft(store, "project", "knowledge/drafts/source.json")

    assert loaded["source_path"] == "source.doc"
    with pytest.raises(ValueError, match="escapes"):
        _load_ir_draft(store, "project", "../outside.json")


def test_bounded_ingestion_supports_markdown_docx_and_image(tmp_path) -> None:
    markdown = tmp_path / "notes.md"
    markdown.write_text("one\ntwo\nthree\nfour\n", encoding="utf-8")
    bounded = read_markdown_lines(markdown, offset=2, limit=2)
    assert bounded.text == "two\nthree"
    assert (bounded.start_line, bounded.end_line) == (2, 3)
    assert read_bounded_source(markdown, offset=3, limit=1).text == "three"

    from docx import Document

    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Paragraph evidence")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    document.save(docx_path)
    docx_read = read_docx_text_tables(docx_path, offset=1, limit=10)
    assert "Paragraph evidence" in docx_read.text
    assert "A | B" in docx_read.text

    from PIL import Image

    image_path = tmp_path / "figure.png"
    Image.new("RGB", (12, 8), color="white").save(image_path)
    image_read = read_image_vision_ocr(image_path)
    assert image_read.image_path.endswith("figure.png")
    assert image_read.metadata and image_read.metadata["width"] == 12


def test_candidate_requires_review_before_live_apply_and_graph_is_undirected(tmp_path) -> None:
    source_root = tmp_path / "sources"
    source_root.mkdir()
    (source_root / "runtime.md").write_text("# Runtime\n\nEvidence text.\n", encoding="utf-8")
    service = KnowledgeService(KnowledgeStore(tmp_path))
    project_id = service.scan(str(source_root), title="Structured Wiki")["project"]["id"]
    service.extract(
        project_id,
        "runtime.md",
        pages=[{
            "type": "concept",
            "title": "Runtime",
            "slug": "runtime",
            "body": (
                "A durable execution runtime coordinates bounded tool calls, keeps state observable, "
                "and preserves source-linked outcomes for later review."
            ),
            "tags": ["runtime"],
            "sources": ["runtime.md"],
        }, {
            "type": "entity",
            "title": "Tools",
            "slug": "tools",
            "body": (
                "A tool boundary exposes bounded operations to the runtime, records each invocation, "
                "and prevents unrelated source material from entering the active context."
            ),
            "tags": ["tools"],
            "sources": ["runtime.md"],
        }],
        relations=[
                {
                    "source": "runtime",
                    "target": "tools",
                    "relation": "uses",
                    "evidence": "The runtime uses tools.",
                },
                {
                    "source": "tools",
                    "target": "runtime",
                    "relation": "related",
                    "evidence": "The tool boundary belongs to the runtime.",
                },
            ],
        evidence=[{
            "source_path": "runtime.md",
            "start_line": 1,
            "end_line": 2,
            "quote": "Evidence text.",
            "extraction_method": "markdown_lines",
        }],
        claims=[{
            "subject": "Runtime",
            "predicate": "uses",
            "object": "Tools",
            "source_path": "runtime.md",
            "evidence": [{"source_path": "runtime.md", "start_line": 1, "quote": "Evidence text."}],
        }],
        review_hints=[{
            "kind": "confirmation",
            "title": "Confirm relation semantics",
            "summary": "Confirm that the runtime-tool relationship is intended.",
            "source_refs": ["runtime.md"],
        }],
    )

    candidate = service.compile_candidate(project_id, reason="review topology")
    changeset_id = candidate["changeset"]["id"]
    project_root = tmp_path / "wikis" / project_id
    assert not (project_root / "wiki" / "index.md").exists()
    assert (project_root / "knowledge" / "candidates" / changeset_id / "wiki" / "index.md").exists()

    review = service.validate_candidate(project_id, changeset_id)
    assert review["passed"] is True, review
    assert any(
        issue["kind"] == "confirmation"
        for issue in review["review"]["issues"]
    )
    stale_page = project_root / "wiki" / "entities" / "obsolete.md"
    stale_page.parent.mkdir(parents=True, exist_ok=True)
    stale_page.write_text("obsolete compiler output", encoding="utf-8")
    applied = service.approve_changeset(project_id, changeset_id)
    assert applied["applied"] is True
    assert (project_root / "wiki" / "index.md").exists()
    assert not stale_page.exists()

    graph = json.loads((project_root / "knowledge" / "graph" / "graph.json").read_text(encoding="utf-8"))
    assert graph["directed"] is False
    assert len(graph["edges"]) == 1
    edge = graph["edges"][0]
    assert {edge["source"], edge["target"]} == {"runtime", "tools"}
    assert all(node["community_id"] and node["community_size"] >= 1 for node in graph["nodes"])
    project = KnowledgeStore(tmp_path).get_project(project_id)
    assert project.review_status == "approved"
    assert project.published_revision_id == applied["revision_id"]


async def test_knowledge_publish_cannot_bypass_human_approval(tmp_path) -> None:
    source_root = tmp_path / "sources"
    source_root.mkdir()
    (source_root / "runtime.md").write_text("# Runtime\n\nEvidence text.\n", encoding="utf-8")
    service = KnowledgeService(KnowledgeStore(tmp_path))
    project_id = service.scan(str(source_root), title="Approval Wiki")["project"]["id"]
    service.extract(
        project_id,
        "runtime.md",
        pages=[{
            "type": "concept",
            "title": "Runtime",
            "slug": "runtime",
            "body": (
                "A durable execution runtime coordinates bounded tool calls, keeps state observable, "
                "and preserves source-linked outcomes for later review."
            ),
            "tags": ["runtime"],
            "sources": ["runtime.md"],
            "evidence": [{
                "source_path": "runtime.md",
                "start_line": 1,
                "end_line": 2,
                "quote": "Evidence text.",
                "extraction_method": "markdown_lines",
            }],
        }],
        evidence=[{
            "source_path": "runtime.md",
            "start_line": 1,
            "end_line": 2,
            "quote": "Evidence text.",
            "extraction_method": "markdown_lines",
        }],
        claims=[{
            "subject": "Runtime",
            "predicate": "supports",
            "object": "bounded tools",
            "source_path": "runtime.md",
            "evidence": [{
                "source_path": "runtime.md",
                "start_line": 1,
                "end_line": 2,
                "quote": "Evidence text.",
            }],
        }],
    )
    changeset_id = service.compile_candidate(project_id)["changeset"]["id"]
    sessions = SessionManager(tmp_path)
    session_key = "websocket:knowledge-approval"
    request = RequestContext(
        channel="websocket",
        chat_id="knowledge-approval",
        session_key=session_key,
        workspace=tmp_path,
    )
    tool = KnowledgePublishTool(str(tmp_path), sessions)

    with request_context(request):
        first = await tool.execute(project_id=project_id, changeset_id=changeset_id)
    assert getattr(first, "is_error", False) is True
    session = sessions.get_or_create(session_key)
    interaction = session.metadata["interaction_request"]
    assert interaction["status"] == "pending"
    assert interaction["_server"]["changeset_id"] == changeset_id
    assert not (tmp_path / "wikis" / project_id / "wiki" / "index.md").exists()

    with request_context(request):
        waiting = await tool.execute(project_id=project_id, changeset_id=changeset_id)
    assert getattr(waiting, "is_error", False) is True
    assert "already waiting" in str(waiting)

    resolve_interaction(
        session.metadata,
        interaction_id=interaction["id"],
        action="apply_once",
        values={},
    )
    sessions.save(session)
    with request_context(request):
        applied = await tool.execute(project_id=project_id, changeset_id=changeset_id)
    payload = json.loads(str(applied))
    assert payload["applied"] is True
    assert (tmp_path / "wikis" / project_id / "wiki" / "index.md").exists()
    metadata, _ = parse_frontmatter(
        (tmp_path / "wikis" / project_id / "wiki" / "concepts" / "runtime.md").read_text(
            encoding="utf-8"
        )
    )
    assert metadata["evidence"][0]["source_path"] == "runtime.md"
    search = KnowledgeSearchTool(str(tmp_path), sessions)
    with request_context(request):
        search_result = await search.execute(query="Runtime", project_id=project_id)
    search_payload = json.loads(str(search_result))
    assert search_payload["claims"][0]["predicate"] == "supports"


async def test_knowledge_publish_honors_auto_apply_policy(tmp_path) -> None:
    source_root = tmp_path / "sources"
    source_root.mkdir()
    (source_root / "runtime.md").write_text("# Runtime\n\nEvidence text.\n", encoding="utf-8")
    service = KnowledgeService(KnowledgeStore(tmp_path))
    project_id = service.scan(str(source_root), title="Auto Wiki")["project"]["id"]
    service.extract(
        project_id,
        "runtime.md",
        pages=[{
            "type": "concept",
            "title": "Runtime",
            "slug": "runtime",
            "body": (
                "A durable execution runtime coordinates bounded tool calls, records observable state, "
                "and preserves source-linked outcomes for later review."
            ),
            "tags": ["runtime"],
            "sources": ["runtime.md"],
        }],
    )
    changeset_id = service.compile_candidate(project_id)["changeset"]["id"]
    sessions = SessionManager(tmp_path)
    session_key = "websocket:knowledge-auto"
    request = RequestContext(
        channel="websocket",
        chat_id="knowledge-auto",
        session_key=session_key,
        workspace=tmp_path,
    )
    tool = KnowledgePublishTool(str(tmp_path), sessions)
    scope = build_workspace_scope(tmp_path, "restricted", execution_policy="auto")
    token = bind_workspace_scope(scope)
    try:
        with request_context(request):
            applied = await tool.execute(project_id=project_id, changeset_id=changeset_id)
    finally:
        reset_workspace_scope(token)

    payload = json.loads(str(applied))
    assert payload["applied"] is True
    assert payload["revision_id"]
    session = sessions.get_or_create(session_key)
    assert "interaction_request" not in session.metadata
    assert (tmp_path / "wikis" / project_id / "wiki" / "index.md").exists()


async def test_knowledge_publish_reads_persisted_auto_policy_without_context_binding(tmp_path) -> None:
    source_root = tmp_path / "sources"
    source_root.mkdir()
    (source_root / "runtime.md").write_text("# Runtime\n\nEvidence text.\n", encoding="utf-8")
    service = KnowledgeService(KnowledgeStore(tmp_path))
    project_id = service.scan(str(source_root), title="Persisted Auto Wiki")["project"]["id"]
    service.extract(
        project_id,
        "runtime.md",
        pages=[{
            "type": "concept",
            "title": "Runtime",
            "slug": "runtime",
            "body": (
                "A durable execution runtime coordinates bounded tool calls, records observable state, "
                "and preserves source-linked outcomes for later review."
            ),
            "tags": ["runtime"],
            "sources": ["runtime.md"],
        }],
    )
    changeset_id = service.compile_candidate(project_id)["changeset"]["id"]
    sessions = SessionManager(tmp_path)
    session_key = "websocket:knowledge-persisted-auto"
    session = sessions.get_or_create(session_key)
    session.metadata["workspace_scope"] = build_workspace_scope(
        tmp_path,
        "restricted",
        execution_policy="auto",
    ).metadata()
    sessions.save(session)
    request = RequestContext(
        channel="websocket",
        chat_id="knowledge-persisted-auto",
        session_key=session_key,
        workspace=tmp_path,
    )
    tool = KnowledgePublishTool(str(tmp_path), sessions)

    # This simulates an interaction-resume/isolated tool call where the
    # AgentLoop ContextVar was not propagated, while session metadata is real.
    with request_context(request):
        applied = await tool.execute(project_id=project_id, changeset_id=changeset_id)

    payload = json.loads(str(applied))
    assert payload["applied"] is True
    assert payload["execution_policy"] == "auto"
    assert "interaction_request" not in sessions.get_or_create(session_key).metadata
