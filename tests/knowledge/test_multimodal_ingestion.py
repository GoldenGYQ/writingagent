from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from nanobot.knowledge.docx_extract import extract_docx_structured
from nanobot.knowledge.ingest import adapter_for_path, read_bounded_source
from nanobot.knowledge.models import KnowledgeChunk
from nanobot.knowledge.normalization import normalize_source
from nanobot.knowledge.service import KnowledgeService
from nanobot.knowledge.store import KnowledgeStore
from nanobot.knowledge.vector_store import LocalVectorStore


def _sample_docx(tmp_path: Path) -> Path:
    image = tmp_path / "invoice.png"
    Image.new("RGB", (120, 60), color="white").save(image)
    source = tmp_path / "evidence.docx"
    document = Document()
    document.add_paragraph("企业资质与履约证据")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "证书"
    table.cell(0, 1).text = "编号"
    table.cell(1, 0).text = "质量体系"
    table.cell(1, 1).text = "Q-001"
    document.add_picture(str(image), width=Inches(1))
    document.save(source)
    return source


def test_doc_and_docx_adapters_are_distinct() -> None:
    assert adapter_for_path("large.doc").name == "legacy_doc_ole"
    assert adapter_for_path("modern.docx").name == "docx_text_tables"


def test_structured_docx_extraction_preserves_assets_and_tables(tmp_path: Path) -> None:
    source = _sample_docx(tmp_path)
    output = tmp_path / "normalized"

    manifest = extract_docx_structured(source, output)

    assert "企业资质与履约证据" in (output / "text.md").read_text(encoding="utf-8")
    assert manifest["tables"][0]["rows"] == 2
    assert len(manifest["images"]) == 1
    assert (output / manifest["images"][0]["path"]).is_file()
    assert json.loads((output / "manifest.json").read_text(encoding="utf-8"))["errors"] == []


def test_docx_bounded_reader_does_not_return_whole_document(tmp_path: Path) -> None:
    source = _sample_docx(tmp_path)
    result = read_bounded_source(source, offset=2, limit=2)
    assert result.adapter == "docx_text_tables"
    assert result.start_line == 2
    assert result.end_line == 3


def test_normalization_uses_bounded_ocr_budget(tmp_path: Path, monkeypatch) -> None:
    source = _sample_docx(tmp_path)
    monkeypatch.setattr(
        "nanobot.knowledge.normalization.ocr_image",
        lambda _path: type("Result", (), {
            "text": "发票号码 123",
            "engine": "test",
            "confidence": 0.9,
            "available": True,
            "error": "",
        })(),
    )

    manifest = normalize_source(source, tmp_path / "result", max_ocr_assets=1)

    assert manifest["status"] == "normalized"
    assert manifest["ocr"][0]["text"] == "发票号码 123"
    assert manifest["ocr"][0]["document_type"] == "invoice"
    assert manifest["image_catalog"] == {
        "jsonl": "image-catalog.jsonl",
        "markdown": "image-catalog.md",
    }
    assert (tmp_path / "result" / "image-catalog.md").is_file()
    assert manifest["adapter"] == "docx_structured"


def test_normalization_resumes_ocr_without_recovering_legacy_assets(tmp_path: Path, monkeypatch) -> None:
    source = tmp_path / "large.doc"
    source.write_bytes(b"legacy-doc-placeholder")
    output = tmp_path / "result"
    calls = {"recover": 0, "ocr": 0}

    monkeypatch.setattr(
        "nanobot.knowledge.normalization.recover_legacy_doc_text",
        lambda _path: ("Recovered tender text", {"format": "ole"}),
    )

    def recover_assets(_source, images, *, max_assets):
        calls["recover"] += 1
        images.mkdir(parents=True, exist_ok=True)
        assets = []
        for index in range(3):
            path = images / f"asset-{index}.png"
            path.write_bytes(b"image")
            assets.append({"id": f"asset-{index}", "path": path.as_posix()})
        return {"assets": assets, "limitations": []}

    def fake_ocr(_path):
        calls["ocr"] += 1
        return type("Result", (), {
            "text": f"OCR {calls['ocr']}",
            "engine": "test",
            "confidence": 0.9,
            "available": True,
            "error": "",
        })()

    monkeypatch.setattr("nanobot.knowledge.normalization.recover_legacy_doc_assets", recover_assets)
    monkeypatch.setattr("nanobot.knowledge.normalization.ocr_image", fake_ocr)

    first = normalize_source(source, output, max_ocr_assets=1)
    second = normalize_source(source, output, max_ocr_assets=1, max_legacy_assets=0)

    assert calls == {"recover": 1, "ocr": 2}
    assert len(first["ocr"]) == 1
    assert len(second["ocr"]) == 2
    assert second["resumed"] is True


def test_legacy_normalization_rejects_resume_only_without_assets(tmp_path: Path) -> None:
    source = tmp_path / "large.doc"
    source.write_bytes(b"legacy-doc-placeholder")

    with pytest.raises(ValueError, match="resume-only"):
        normalize_source(
            source,
            tmp_path / "result",
            max_ocr_assets=0,
            max_legacy_assets=0,
        )


def test_fastembed_backend_has_stable_index_contract(tmp_path: Path, monkeypatch) -> None:
    class FakeEmbedding:
        def __init__(self, *, model_name: str) -> None:
            self.model_name = model_name

        def embed(self, texts):
            for text in texts:
                yield [1.0, float("履约" in text), 0.0]

    monkeypatch.setattr("fastembed.TextEmbedding", FakeEmbedding)
    chunk = KnowledgeChunk(
        id="chunk",
        document_id="doc",
        project_id="project",
        path="concepts/guarantee.md",
        title="履约保证",
        page_type="concept",
        text="履约保证金",
        start_line=1,
        end_line=1,
    )
    store = LocalVectorStore(
        tmp_path / "vectors.json",
        backend="fastembed",
        model_name="BAAI/bge-small-zh-v1.5",
    )
    assert store.build([chunk]) == 1
    store.save()

    restored = LocalVectorStore(
        tmp_path / "vectors.json",
        backend="fastembed",
        model_name="BAAI/bge-small-zh-v1.5",
    )
    assert restored.load() == 1
    assert restored.search("履约")


def test_scan_streams_raw_sources_without_path_read_bytes(tmp_path: Path, monkeypatch) -> None:
    workspace = tmp_path / "workspace"
    source = workspace / "sources"
    source.mkdir(parents=True)
    (source / "large.doc").write_bytes(b"legacy-doc-placeholder" * 10_000)

    monkeypatch.setattr(
        Path,
        "read_bytes",
        lambda _self: (_ for _ in ()).throw(AssertionError("read_bytes must not be used")),
    )
    result = KnowledgeService(KnowledgeStore(workspace)).scan(str(source))

    assert result["files"] == 1
    assert result["project"]["sources"][0]["sha256"]
