"""Observable source-ingestion adapter contracts.

The Knowledge Runtime keeps ingestion separate from extraction and publishing.
This module only describes how an Agent should read a raw source; the raw bytes
are still mirrored by :class:`KnowledgeService` and the model produces the
typed IR through ``knowledge_extract``.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, cast


@dataclass(frozen=True)
class IngestionAdapterSpec:
    """A bounded reader contract exposed in the source manifest."""

    name: str
    kind: str
    extensions: tuple[str, ...]
    extraction_mode: str
    instruction: str
    requires_vision: bool = False

    def metadata(self) -> dict[str, Any]:
        return {
            "ingestion_adapter": self.name,
            "extraction_mode": self.extraction_mode,
            "instruction": self.instruction,
            "requires_vision": self.requires_vision,
            "bounded_read": {
                "mode": self.extraction_mode,
                "instruction": self.instruction,
            },
        }


_TEXT = IngestionAdapterSpec(
    name="text_lines",
    kind="text",
    extensions=(
        ".md", ".markdown", ".txt", ".rst", ".py", ".js", ".ts", ".tsx", ".jsx",
        ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".tex", ".csv",
    ),
    extraction_mode="line_bounded",
    instruction="Use read_file with offset/limit and preserve source line anchors.",
)
_MARKDOWN = IngestionAdapterSpec(
    name="markdown_lines",
    kind="markdown",
    extensions=(".md", ".markdown"),
    extraction_mode="line_bounded",
    instruction="Use read_file with offset/limit and preserve Markdown headings and source line anchors.",
)
_DOCX = IngestionAdapterSpec(
    name="docx_text_tables",
    kind="docx",
    extensions=(".docx",),
    extraction_mode="paragraph_bounded",
    instruction=(
        "Use read_source_range with start_line/end_line over paragraphs and tables; "
        "do not inject the entire DOCX into context."
    ),
)
_DOC = IngestionAdapterSpec(
    name="legacy_doc_ole",
    kind="doc",
    extensions=(".doc",),
    extraction_mode="recovery_and_asset_bounded",
    instruction=(
        "Recover bounded text from the WordDocument stream and normalize embedded assets separately; "
        "never load the complete OLE Data stream or document into model context."
    ),
    requires_vision=True,
)
_PDF = IngestionAdapterSpec(
    name="pdf_pages",
    kind="pdf",
    extensions=(".pdf",),
    extraction_mode="page_bounded",
    instruction="Use read_file with a bounded pages range; do not inject the whole PDF into context.",
)
_HTML = IngestionAdapterSpec(
    name="html_text",
    kind="html",
    extensions=(".html", ".htm"),
    extraction_mode="text_bounded",
    instruction="Read the saved HTML as source text and preserve its local path as evidence.",
)
_IMAGE = IngestionAdapterSpec(
    name="vision_ocr",
    kind="image",
    extensions=(".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"),
    extraction_mode="vision_or_ocr",
    instruction="Use vision/OCR to extract only the needed regions and record the image path as evidence.",
    requires_vision=True,
)

_BY_EXTENSION = {
    extension: spec
    for spec in (_TEXT, _MARKDOWN, _DOCX, _DOC, _PDF, _HTML, _IMAGE)
    for extension in spec.extensions
}


@dataclass(frozen=True)
class BoundedSourceRead:
    """A bounded extraction result suitable for an Agent evidence payload."""

    source_path: str
    adapter: str
    text: str = ""
    start_line: int | None = None
    end_line: int | None = None
    start_page: int | None = None
    end_page: int | None = None
    image_path: str = ""
    metadata: dict[str, Any] | None = None

    def to_dict(self) -> dict[str, Any]:
        value: dict[str, Any] = {
            "source_path": self.source_path,
            "adapter": self.adapter,
            "text": self.text,
            "start_line": self.start_line,
            "end_line": self.end_line,
            "start_page": self.start_page,
            "end_page": self.end_page,
            "image_path": self.image_path,
        }
        if self.metadata:
            value["metadata"] = dict(self.metadata)
        return value


class IngestionDependencyError(RuntimeError):
    """Raised when an optional local reader is not installed."""


def _bounded_lines(lines: Iterable[str], *, offset: int = 1, limit: int = 240) -> tuple[str, int | None, int | None]:
    start = max(1, offset)
    end_limit = start + max(1, min(limit, 2_000)) - 1
    selected: list[str] = []
    for line_number, line in enumerate(lines, start=1):
        if line_number < start:
            continue
        if line_number > end_limit:
            break
        selected.append(line.rstrip("\r\n"))
    if not selected:
        return "", None, None
    return "\n".join(selected), start, start + len(selected) - 1


def read_text_lines(path: str | Path, *, offset: int = 1, limit: int = 240) -> BoundedSourceRead:
    source = Path(path)
    text = source.read_text(encoding="utf-8", errors="replace")
    value, start, end = _bounded_lines(text.splitlines(), offset=offset, limit=limit)
    return BoundedSourceRead(source.as_posix(), "text_lines", value, start, end)


def read_markdown_lines(path: str | Path, *, offset: int = 1, limit: int = 240) -> BoundedSourceRead:
    result = read_text_lines(path, offset=offset, limit=limit)
    return BoundedSourceRead(
        result.source_path,
        "markdown_lines",
        result.text,
        result.start_line,
        result.end_line,
        metadata={"preserve_headings": True},
    )


def _docx_lines(path: Path) -> Iterable[str]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised in dependency-gap environments
        raise IngestionDependencyError("python-docx is required for DOCX ingestion") from exc
    document = Document(str(path))
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            yield value
    for table_index, table in enumerate(document.tables, start=1):
        yield f"[Table {table_index}]"
        for row in table.rows:
            yield " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)


def read_docx_text_tables(path: str | Path, *, offset: int = 1, limit: int = 240) -> BoundedSourceRead:
    source = Path(path)
    value, start, end = _bounded_lines(_docx_lines(source), offset=offset, limit=limit)
    return BoundedSourceRead(source.as_posix(), "docx_text_tables", value, start, end)


def read_legacy_doc_ole(path: str | Path, *, offset: int = 1, limit: int = 240) -> BoundedSourceRead:
    from nanobot.knowledge.legacy_doc import recover_legacy_doc_text

    source = Path(path)
    text, metadata = recover_legacy_doc_text(source)
    value, start, end = _bounded_lines(text.splitlines(), offset=offset, limit=limit)
    return BoundedSourceRead(source.as_posix(), "legacy_doc_ole", value, start, end, metadata=metadata)


def read_pdf_pages(
    path: str | Path,
    *,
    start_page: int = 1,
    end_page: int | None = None,
    max_chars: int = 40_000,
) -> BoundedSourceRead:
    source = Path(path)
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - exercised in dependency-gap environments
        raise IngestionDependencyError("pypdf is required for PDF ingestion") from exc
    reader = PdfReader(str(source))
    first = max(1, start_page)
    last = min(len(reader.pages), end_page or first + 4)
    if last < first:
        return BoundedSourceRead(source.as_posix(), "pdf_pages", start_page=first, end_page=last)
    chunks: list[str] = []
    ocr_pages: list[int] = []
    for page_number in range(first, last + 1):
        text = reader.pages[page_number - 1].extract_text() or ""
        if len(text.strip()) < 24:
            ocr_text = _ocr_pdf_page(source, page_number)
            if ocr_text:
                text = ocr_text
                ocr_pages.append(page_number)
        chunks.append(f"[Page {page_number}]\n{text.strip()}")
        if sum(len(item) for item in chunks) >= max(1_000, max_chars):
            break
    combined = "\n\n".join(chunks)[:max(1_000, max_chars)]
    actual_last = first + len(chunks) - 1
    return BoundedSourceRead(
        source.as_posix(),
        "pdf_pages",
        combined,
        start_page=first,
        end_page=actual_last,
        metadata={"total_pages": len(reader.pages), "ocr_pages": ocr_pages},
    )


def _ocr_pdf_page(path: Path, page_number: int) -> str:
    """Render and OCR one page when the PDF text layer is empty."""

    import tempfile

    try:
        import fitz  # pyright: ignore[reportMissingTypeStubs]
    except ImportError:
        return ""
    from nanobot.knowledge.ocr import ocr_image

    document_factory = cast(Any, fitz.open)
    matrix_factory = cast(Any, fitz.Matrix)
    with document_factory(path) as document:
        page = document.load_page(page_number - 1)
        pixmap = page.get_pixmap(matrix=matrix_factory(2.0, 2.0), alpha=False)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temporary:
            temporary_path = Path(temporary.name)
        try:
            pixmap.save(temporary_path)
            return ocr_image(temporary_path).text
        finally:
            temporary_path.unlink(missing_ok=True)


def read_image_vision_ocr(path: str | Path) -> BoundedSourceRead:
    source = Path(path)
    metadata: dict[str, Any] = {}
    try:
        from PIL import Image

        with Image.open(source) as image:
            metadata.update({"width": image.width, "height": image.height, "format": image.format or ""})
    except ImportError:  # pragma: no cover - Pillow is optional
        metadata["image_reader"] = "unavailable"
    except OSError as exc:
        metadata["image_error"] = str(exc)[:500]
    from nanobot.knowledge.ocr import ocr_image

    result = ocr_image(source)
    ocr_text = result.text
    metadata.update(result.metadata())
    return BoundedSourceRead(
        source.as_posix(),
        "vision_ocr",
        ocr_text[:40_000],
        image_path=source.as_posix(),
        metadata=metadata,
    )


def read_html_text(path: str | Path, *, offset: int = 1, limit: int = 240) -> BoundedSourceRead:
    source = Path(path)
    html = source.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup

        text = BeautifulSoup(html, "html.parser").get_text("\n")
    except ImportError:  # pragma: no cover - exercised in dependency-gap environments
        import re

        text = re.sub(r"<[^>]+>", " ", html)
    value, start, end = _bounded_lines(text.splitlines(), offset=offset, limit=limit)
    return BoundedSourceRead(source.as_posix(), "html_text", value, start, end)


def read_bounded_source(path: str | Path, **kwargs: Any) -> BoundedSourceRead:
    """Dispatch to a bounded reader without reading unrelated source content."""
    adapter = adapter_for_path(path)
    if adapter is None:
        raise ValueError(f"unsupported knowledge source: {path}")
    if adapter.name == "markdown_lines":
        return read_markdown_lines(path, **kwargs)
    if adapter.name == "text_lines":
        return read_text_lines(path, **kwargs)
    if adapter.name == "docx_text_tables":
        return read_docx_text_tables(path, **kwargs)
    if adapter.name == "legacy_doc_ole":
        return read_legacy_doc_ole(path, **kwargs)
    if adapter.name == "pdf_pages":
        return read_pdf_pages(path, **kwargs)
    if adapter.name == "vision_ocr":
        return read_image_vision_ocr(path)
    if adapter.name == "html_text":
        return read_html_text(path, **kwargs)
    raise ValueError(f"no reader for adapter: {adapter.name}")


def adapter_for_path(path: str | Path) -> IngestionAdapterSpec | None:
    """Return the reader contract for a path suffix, if supported."""
    return _BY_EXTENSION.get(Path(path).suffix.lower())


def supported_source_suffixes() -> frozenset[str]:
    """Return the suffixes that ``knowledge_scan`` may mirror into ``raw``."""
    return frozenset(_BY_EXTENSION)
